from __future__ import annotations

import io

from fastapi import APIRouter, File, Form, Response, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from ..api.schemas.asr import ASRTranscriptionResponse
from ..config import get_settings
from ..services.asr_cloud_service import local_asr_health, local_asr_languages, transcribe_audio

router = APIRouter(prefix="/asr", tags=["asr"])


@router.post("/transcribe", response_model=ASRTranscriptionResponse)
async def transcribe(
    audio: UploadFile = File(...),
    provider: str | None = Form(default=None),
    language: str | None = Form(default=None),
    speakers: int | None = Form(default=None),
    diarize: bool = Form(default=True),
) -> ASRTranscriptionResponse:
    return await transcribe_audio(
        settings=get_settings(),
        audio=audio,
        provider=provider,
        language=language,
        num_speakers=speakers,
        diarize=diarize,
    )


@router.get("/local/health")
async def local_health() -> dict:
    return await local_asr_health(get_settings())


@router.get("/local/languages")
async def local_languages() -> dict:
    return await local_asr_languages(get_settings())


class ExportSegment(BaseModel):
    speaker: str
    time: str = ""
    text: str


class ExportDocxRequest(BaseModel):
    segments: list[ExportSegment] = Field(default_factory=list)
    meta: dict[str, str | int | float] = Field(default_factory=dict)


@router.post("/export/docx")
async def export_docx(payload: ExportDocxRequest) -> Response:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading("Transkripsiya", level=0)

    info: list[str] = []
    if payload.meta.get("language"):
        info.append(f"Til: {payload.meta['language']}")
    if payload.meta.get("duration"):
        info.append(f"Davomiyligi: {payload.meta['duration']}")
    if payload.meta.get("speakersCount"):
        info.append(f"Spikerlar: {payload.meta['speakersCount']}")
    if info:
        p = doc.add_paragraph()
        run = p.add_run(" · ".join(info))
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x46, 0x45, 0x55)

    for segment in payload.segments:
        head = doc.add_paragraph()
        speaker_run = head.add_run(segment.speaker)
        speaker_run.bold = True
        speaker_run.font.size = Pt(11)
        speaker_run.font.color.rgb = RGBColor(0x35, 0x25, 0xCD)
        if segment.time:
            time_run = head.add_run(f"   {segment.time}")
            time_run.italic = True
            time_run.font.size = Pt(9)
            time_run.font.color.rgb = RGBColor(0x46, 0x45, 0x55)
        doc.add_paragraph(segment.text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="transkripsiya.docx"'},
    )
